import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches
import os


doc = Document()
doc.add_heading('3D Printer Usage Report: 2023 vs 2024', 0)

os.makedirs("temp_figs", exist_ok=True)

plt.style.use('ggplot')
months = ['Jan','Feb','Mar','Apr','May','Jun','Jul','Aug','Sep','Oct','Nov','Dec']

def save_and_insert(fig, filename, doc, caption=None):
    fig_path = f"temp_figs/{filename}.png"
    fig.savefig(fig_path, bbox_inches='tight')
    doc.add_picture(fig_path, width=Inches(6))
    if caption:
        doc.add_paragraph(caption, style='Caption')

# === 1: Unique Users on Jan 1 ===
years = ['2023', '2024']
users_jan1 = [1032, 1300]
fig, ax = plt.subplots()
ax.bar(years, users_jan1, color=['blue', 'green'])
ax.set_title('Unique Users on Jan 1')
ax.set_ylabel('Users')
plt.subplots_adjust(bottom=0.35)
table = plt.table(cellText=[['2023', '1032'], ['2024', '1300']],
                  colLabels=['Year', 'Users'],
                  cellLoc='center',
                  bbox=[0.3, -0.35, 0.4, 0.2])
table.scale(1, 1.2)
# plt.show()
fig = plt.gcf()
save_and_insert(fig, 'unique_users_jan1', doc, caption='')
plt.close(fig)

# === 2: Unique Users on Dec 31 ===
fig, ax = plt.subplots()
plt.bar(['2023', '2024'], [1300, 1591], color=['blue', 'green'])
plt.title('Unique Users on Dec 31')
plt.ylabel('Users')
plt.subplots_adjust(bottom=0.3)
table = plt.table(cellText=[['2023', '1300'], ['2024', '1591']],
                  colLabels=['Year', 'Users'],
                  cellLoc='center',
                  bbox=[0.3, -0.35, 0.4, 0.2])
table.scale(1, 1.2)
# plt.show()
fig = plt.gcf()
save_and_insert(fig, 'unique_users_dec31', doc, caption='')
plt.close(fig)




# === 3: User Categories ===
labels = ['Active', 'Inactive', 'Multiple Jobs', 'Single Job']
data_2023 = [227, 1073, 136, 91]
data_2024 = [267, 1324, 174, 93]
x = range(len(labels))
fig, ax = plt.subplots()
bar_width = 0.4
plt.bar(x, data_2023, width=bar_width, label='2023', align='center')
plt.bar([p + bar_width for p in x], data_2024, width=bar_width, label='2024', align='center')
plt.xticks([p + bar_width / 2 for p in x], labels)
plt.title('User Categories')
plt.legend()
plt.subplots_adjust(bottom=0.4)
table = plt.table(cellText=[['Active', 'Inactive', 'Multiple Jobs', 'Single Job'],
                            ['227', '1073', '136', '91'],
                            ['267', '1324', '174', '93']],
                  rowLabels=['Year', '2023', '2024'],
                  cellLoc='center',
                  bbox=[0.0, -0.45, 1.0, 0.3])
table.scale(1, 1.2)
# plt.show()
fig = plt.gcf()
save_and_insert(fig, 'user_categories', doc, caption='')
plt.close(fig)

# === 4: New Users ===
fig, ax = plt.subplots()
plt.bar(['2023', '2024'], [268, 291], color=['blue', 'green'])
plt.title('New/First-Time Users')
plt.ylabel('Users')
plt.subplots_adjust(bottom=0.3)
table = plt.table(cellText=[['2023', '268'], ['2024', '291']],
                  colLabels=['Year', 'New Users'],
                  cellLoc='center',
                  bbox=[0.3, -0.35, 0.4, 0.2])
table.scale(1, 1.2)
# plt.show()
fig = plt.gcf()
save_and_insert(fig, 'new_users', doc, caption='')
plt.close(fig)

# === 5: Print Jobs Submitted ===
jobs_2023 = [89, 82, 78, 75, 34, 67, 96, 137, 92, 92, 166, 68]
jobs_2024 = [154, 106, 103, 81, 64, 72, 78, 80, 69, 108, 112, 156]

# Branch-specific data
jobs_central_2023 = [89, 82, 77, 75, 34, 52, 78, 109, 46, 72, 96, 36]
jobs_sunrise_2023 = [0, 0, 1, 0, 0, 15, 18, 28, 46, 20, 70, 32]

jobs_central_2024 = [90, 61, 70, 42, 41, 39, 47, 44, 41, 64, 59, 68]
jobs_sunrise_2024 = [64, 45, 33, 39, 23, 33, 31, 36, 28, 44, 53, 88]

# Simplified comparison bars

# Totals
total_2023 = sum(jobs_2023)
total_2024 = sum(jobs_2024)
central_2023 = sum(jobs_central_2023)
central_2024 = sum(jobs_central_2024)
sunrise_2023 = sum(jobs_sunrise_2023)
sunrise_2024 = sum(jobs_sunrise_2024)

# Plot 1: Total Jobs Comparison
plt.figure(figsize=(8, 5))
plt.bar(['2023', '2024'], [total_2023, total_2024], color=['blue', 'green'])
plt.title('Total Print Jobs Submitted: 2023 vs 2024')
plt.ylabel('Jobs')

# Value labels
plt.text(0, total_2023 + 10, str(total_2023), ha='center', va='bottom', fontsize=10)
plt.text(1, total_2024 + 10, str(total_2024), ha='center', va='bottom', fontsize=10)

table = plt.table(cellText=[['2023', total_2023], ['2024', total_2024]],
                  colLabels=['Year', 'Submitted Jobs'],
                  cellLoc='center',
                  bbox=[0.3, -0.35, 0.4, 0.2])
table.scale(1, 1.2)

plt.tight_layout()
# plt.show()
fig = plt.gcf()
save_and_insert(fig, 'total_printed', doc, caption='')
plt.close(fig)

# Plot 2: Branch Breakdown
plt.figure(figsize=(8, 5))
labels = ['Central 2023', 'Central 2024', 'Sunrise 2023', 'Sunrise 2024']
values = [central_2023, central_2024, sunrise_2023, sunrise_2024]
colors = ['blue', 'green', 'orange', 'purple']

plt.bar(labels, values, color=colors)
plt.title('Branch-wise Total Print Jobs Submitted: 2023 vs 2024')
plt.ylabel('Jobs')

# Value labels
for i, v in enumerate(values):
    plt.text(i, v + 10, str(v), ha='center', va='bottom', fontsize=10)
    
table = plt.table(cellText=[
                    ['2023', central_2023, sunrise_2023, total_2023],
                    ['2024', central_2024, sunrise_2024, total_2024]
                 ],
                 colLabels=['Year', 'Central', 'Sunrise', 'Total'],
                 cellLoc='center',
                 bbox=[0.2, -0.3, 0.6, 0.2])
table.scale(1, 1.2)

plt.tight_layout()
# plt.show()
fig = plt.gcf()
save_and_insert(fig, 'total_printed_branch', doc, caption='')
plt.close(fig)


#########MONTHLY BREAKDOWN############
# Create a grouped bar chart with monthly breakdown and include a table

# Prepare data
month_labels = ['Jan', 'Feb', 'Mar', 'Apr', 'May', 'Jun',
                'Jul', 'Aug', 'Sep', 'Oct', 'Nov', 'Dec']
x = list(range(12))
bar_width = 0.35
x_2023 = [i - bar_width / 2 for i in x]
x_2024 = [i + bar_width / 2 for i in x]

# Create figure
plt.figure(figsize=(14, 7))

# Bar plots for each year
plt.bar(x_2023, jobs_2023, width=bar_width, label='2023', color='blue')
plt.bar(x_2024, jobs_2024, width=bar_width, label='2024', color='green')

# Add value labels on top of bars
for i in range(12):
    plt.text(x_2023[i], jobs_2023[i] + 2, str(jobs_2023[i]), ha='center', va='bottom', fontsize=8, color='blue')
    plt.text(x_2024[i], jobs_2024[i] + 2, str(jobs_2024[i]), ha='center', va='bottom', fontsize=8, color='green')

# Labels and formatting
plt.xticks(x, month_labels)
plt.ylabel('Number of Jobs')
plt.title('Total Monthly Print Jobs Submitted Breakdown: 2023 vs 2024')
plt.legend()
plt.tight_layout()

# Table data: each row is [month, 2023 jobs, 2024 jobs]
table_data = [[month_labels[i], jobs_2023[i], jobs_2024[i]] for i in range(12)]

# Add total row
table_data.append(['TOTAL', sum(jobs_2023), sum(jobs_2024)])

# Add the table
table = plt.table(cellText=table_data,
                  colLabels=['Month', '2023', '2024'],
                  cellLoc='center',
                  bbox=[0.1, -0.45, 0.8, 0.4])
table.scale(1, 1.2)

# plt.show()
fig = plt.gcf()
save_and_insert(fig, 'total_monthly', doc, caption='')
plt.close(fig)

#####################################################################################

# === 6: Print Jobs Printed ===
printed_2023 = [65, 73, 57, 57, 26, 38, 44, 113, 68, 52, 123, 56]
printed_2024 = [103, 84, 57, 29, 44, 61, 60, 55, 48, 65, 80, 95]
printed_central_2024 = [54, 39, 45, 19, 33, 32, 36, 28, 30, 41, 39, 37]
printed_sunrise_2024 = [49, 45, 12, 10, 11, 29, 24, 27, 18, 24, 41, 58]
printed_central_2023 = [65, 73, 56, 57, 26, 27, 31, 91, 33, 39, 76, 28]
printed_sunrise_2023 = [0, 0, 1, 0, 0, 11, 13, 22, 35, 13, 47, 28]

# Totals
total_printed_2023 = sum(printed_2023)
total_printed_2024 = sum(printed_2024)
central_printed_2023 = sum(printed_central_2023)
central_printed_2024 = sum(printed_central_2024)
sunrise_printed_2023 = sum(printed_sunrise_2023)
sunrise_printed_2024 = sum(printed_sunrise_2024)

# --- Plot 1: Total print jobs printed ---
plt.figure(figsize=(8, 5))
plt.bar(['2023', '2024'], [total_printed_2023, total_printed_2024], color=['blue', 'green'])
plt.title('Total Print Jobs Printed: 2023 vs 2024')
plt.ylabel('Jobs')

# Value labels
plt.text(0, total_printed_2023 + 10, str(total_printed_2023), ha='center', va='bottom', fontsize=10)
plt.text(1, total_printed_2024 + 10, str(total_printed_2024), ha='center', va='bottom', fontsize=10)

# Table below
table = plt.table(cellText=[['2023', total_printed_2023], ['2024', total_printed_2024]],
                  colLabels=['Year', 'Printed Jobs'],
                  cellLoc='center',
                  bbox=[0.3, -0.35, 0.4, 0.2])
table.scale(1, 1.2)

plt.tight_layout()
# plt.show()
fig = plt.gcf()
save_and_insert(fig, 'total_printed', doc, caption='')
plt.close(fig)

# --- Plot 2: Clustered Branch for Printed Jobs ---
plt.figure(figsize=(8, 5))
labels = ['Central 2023', 'Central 2024', 'Sunrise 2023', 'Sunrise 2024']
values = [central_printed_2023, central_printed_2024, sunrise_printed_2023, sunrise_printed_2024]
colors = ['blue', 'green', 'orange', 'purple']

plt.bar(labels, values, color=colors)
plt.title('Branch-wise Print Jobs Printed: 2023 vs 2024')
plt.ylabel('Jobs')

# Value labels
for i, v in enumerate(values):
    plt.text(i, v + 10, str(v), ha='center', va='bottom', fontsize=10)

# Corrected table data (numeric totals only)
table = plt.table(cellText=[
                    ['2023', central_printed_2023, sunrise_printed_2023, total_printed_2023],
                    ['2024', central_printed_2024, sunrise_printed_2024, total_printed_2024]
                 ],
                 colLabels=['Year', 'Central', 'Sunrise', 'Total'],
                 cellLoc='center',
                 bbox=[0.2, -0.3, 0.6, 0.2])
table.scale(1, 1.2)

plt.tight_layout()
# plt.show()
fig = plt.gcf()
save_and_insert(fig, 'total_printed_branch', doc, caption='')
plt.close(fig)

###############################################################################
############## MONTHLY BREAKDOWN FOR PRINTED JOBS 2023 v 2024 #############################

plt.figure(figsize=(14, 7))

# Bar plots for each year
plt.bar(x_2023, printed_2023, width=bar_width, label='2023', color='blue')
plt.bar(x_2024, printed_2024, width=bar_width, label='2024', color='green')

# Add value labels on top of bars
for i in range(12):
    plt.text(x_2023[i], printed_2023[i] + 2, str(printed_2023[i]), ha='center', va='bottom', fontsize=8, color='blue')
    plt.text(x_2024[i], printed_2024[i] + 2, str(printed_2024[i]), ha='center', va='bottom', fontsize=8, color='green')\

# Labels and formatting
plt.xticks(x, month_labels)
plt.ylabel('Number of Jobs Printed')
plt.title('Monthly Print Jobs Printed Breakdown: 2023 vs 2024')
plt.legend()
plt.tight_layout()

table_data = [[month_labels[i], printed_2023[i], printed_2024[i]] for i in range(12)]

# Add total row
table_data.append(['TOTAL', sum(printed_2023), sum(printed_2024)])

# Add the table
table = plt.table(cellText=table_data,
                  colLabels=['Month', '2023', '2024'],
                  cellLoc='center',
                  bbox=[0.1, -0.45, 0.8, 0.4])
table.scale(1, 1.2)

# plt.show()
fig = plt.gcf()
save_and_insert(fig, 'total_printed_monthly', doc, caption='')
plt.close(fig)



###################################################################################################################################################
####################################################################################################################################################

# === 7: $ Value of Jobs Printed ===
value_2023 = [229.6, 167.7, 263, 246.4, 127.2, 153.1, 182.35, 222.6, 283.7, 272.2, 366.63, 227.1]
value_2024 = [451.16, 390, 249.5, 110.1, 262, 222.6, 230.7, 210.2, 240.28, 358, 434.7, 495.4]
value_central_2024 = [283.16, 188, 196.2, 64.8, 204, 139.6, 158.1, 104.4, 171.18, 211.2, 211.6, 234.9]
value_sunrise_2024 = [168, 202, 53.3, 45.3, 58, 83, 72.6, 105.8, 69.1, 146.8, 223.1, 260.5]
value_central_2023 = [229.6, 167.7, 260.4, 246.4, 127.2, 105.4, 149.5, 149.5, 109.1, 182.6, 147.07, 154.6]
value_sunrise_2023 = [0, 0, 2.6, 0, 0, 47.7, 32.85, 73.1, 174.6, 89.6, 219.56, 72.5]

# Totals
total_value_2023 = sum(value_2023)
total_value_2024 = sum(value_2024)
central_value_2023 = sum(value_central_2023)
central_value_2024 = sum(value_central_2024)
sunrise_value_2023 = sum(value_sunrise_2023)
sunrise_value_2024 = sum(value_sunrise_2024)


#PLOT 1: Total value for 2023 and 2024

plt.figure(figsize=(8, 5))
plt.bar(['2023', '2024'], [total_value_2023, total_value_2024], color=['blue', 'green'])
plt.title('Total Value of Jobs Printed: 2023 vs 2024')
plt.ylabel('Value ($)')
plt.text(0, total_value_2023 + 20, f"${total_value_2023:,.2f}", ha='center', va='bottom', fontsize=10)
plt.text(1, total_value_2024 + 20, f"${total_value_2024:,.2f}", ha='center', va='bottom', fontsize=10)
table = plt.table(cellText=[
                    ['2023', f"${total_value_2023:,.2f}"],
                    ['2024', f"${total_value_2024:,.2f}"]
                 ],
                 colLabels=['Year', 'Value of Print Jobs'],
                 cellLoc='center',
                 bbox=[0.3, -0.35, 0.4, 0.2])
table.scale(1, 1.2)
plt.tight_layout()
# plt.show()

fig = plt.gcf()
save_and_insert(fig, 'total_value', doc, caption='')
plt.close(fig)

# --- Plot 2: Clustered Branch-wise ---
plt.figure(figsize=(8, 5))
labels = ['Central 2023', 'Central 2024', 'Sunrise 2023', 'Sunrise 2024']
values = [central_value_2023, central_value_2024, sunrise_value_2023, sunrise_value_2024]
colors = ['blue', 'green', 'orange', 'purple']
plt.bar(labels, values, color=colors)
plt.title('Branch-wise Print Jobs Value: 2023 vs 2024')
plt.ylabel('Value ($)')
for i, v in enumerate(values):
    plt.text(i, v + 10, f"${v:,.2f}", ha='center', va='bottom', fontsize=10)
table = plt.table(cellText=[
                    ['2023', f"${central_value_2023:,.2f}", f"${sunrise_value_2023:,.2f}", f"${total_value_2023:,.2f}"],
                    ['2024', f"${central_value_2024:,.2f}", f"${sunrise_value_2024:,.2f}", f"${total_value_2024:,.2f}"]
                 ],
                 colLabels=['Year', 'Central', 'Sunrise', 'Total'],
                 cellLoc='center',
                 bbox=[0.2, -0.3, 0.6, 0.2])
table.scale(1, 1.2)
plt.tight_layout()
# plt.show()
fig = plt.gcf()
save_and_insert(fig, 'total_value_branches', doc, caption='')
plt.close(fig)

# --- Plot 3: Monthly Breakdown ---
plt.figure(figsize=(14, 7))
plt.bar(x_2023, value_2023, width=bar_width, label='2023', color='blue')
plt.bar(x_2024, value_2024, width=bar_width, label='2024', color='green')
for i in range(12):
    plt.text(x_2023[i], value_2023[i] + 2, f"${value_2023[i]:.2f}", ha='center', va='bottom', fontsize=8, color='blue')
    plt.text(x_2024[i], value_2024[i] + 2, f"${value_2024[i]:.2f}", ha='center', va='bottom', fontsize=8, color='green')
plt.xticks(x, month_labels)
plt.ylabel('Value of Jobs Printed ($)')
plt.title('Monthly Breakdown for Value of Jobs Printed: 2023 vs 2024')
plt.legend()
plt.tight_layout()
table_data = [[month_labels[i], f"${value_2023[i]:.2f}", f"${value_2024[i]:.2f}"] for i in range(12)]
table_data.append(['TOTAL', f"${sum(value_2023):,.2f}", f"${sum(value_2024):,.2f}"])
table = plt.table(cellText=table_data,
                  colLabels=['Month', '2023', '2024'],
                  cellLoc='center',
                  bbox=[0.1, -0.45, 0.8, 0.4])
table.scale(1, 1.2)
# plt.show()

fig = plt.gcf()
save_and_insert(fig, 'total_value_monthly', doc, caption='')
plt.close(fig)




# === 8: Actual Value Paid ===

actual_value_2023 = 2741.58
actual_value_2024 = 3511.94
actual_value_central_2023 = 2029.07
actual_value_central_2024 = 1292.53
actual_value_sunrise_2023 = 2219.41
actual_value_sunrise_2024 = 712.51
plt.figure()
plt.bar(['2023', '2024'], [actual_value_2023, actual_value_2024], color=['blue', 'green'])
plt.title('$ Value of Print Jobs Actually Paid For')
plt.ylabel('Dollars ($)')
table = plt.table(cellText=[
                    ['2023', f"${actual_value_2023:,.2f}"],
                    ['2024', f"${actual_value_2024:,.2f}"]
                 ],
                 colLabels=['Year', 'Value Actually Paid For'],
                 cellLoc='center',
                 bbox=[0.2, -0.3, 0.6, 0.2])
table.scale(1, 1.2)
plt.tight_layout()
# plt.show()

fig = plt.gcf()
save_and_insert(fig, 'total_actual', doc, caption='')
plt.close(fig)

######### CLUSTERED BY BRANCH ################



plt.figure(figsize=(8, 5))
labels = ['Central 2023', 'Central 2024', 'Sunrise 2023', 'Sunrise 2024']
values = [actual_value_central_2023, actual_value_central_2024, actual_value_sunrise_2023, actual_value_sunrise_2024]
colors = ['blue', 'green', 'orange', 'purple']
plt.bar(labels, values, color=colors)
plt.title('Branch-wise Value of Print Jobs Actually Paid For: 2023 vs 2024')
plt.ylabel('Value ($)')
for i, v in enumerate(values):
    plt.text(i, v + 10, f"${v:,.2f}", ha='center', va='bottom', fontsize=10)
table = plt.table(cellText=[
                    ['2023', f"${actual_value_central_2023:,.2f}", f"${actual_value_sunrise_2023:,.2f}", f"${actual_value_2023:,.2f}"],
                    ['2024', f"${actual_value_central_2023:,.2f}", f"${actual_value_sunrise_2024:,.2f}", f"${actual_value_2024:,.2f}"]
                 ],
                 colLabels=['Year', 'Central', 'Sunrise', 'Total'],
                 cellLoc='center',
                 bbox=[0.2, -0.3, 0.6, 0.2])
table.scale(1, 1.2)
plt.tight_layout()
# plt.show()


fig = plt.gcf()
save_and_insert(fig, 'total_actual_branches', doc, caption='')
plt.close(fig)



# === 9: Usage Time in Hours ===


printhours_2023 = 142
printhours_2024 = 163


printdays_2023 = printhours_2023 / 24
printdays_2024 = printhours_2024 / 24


plt.figure()
plt.bar(['2023', '2024'], [printhours_2023, printhours_2024], color=['blue', 'green'])
plt.title('Total Time to Print (Approval to FInished in Hours)')
plt.ylabel('Hours')


plt.text(0, printhours_2023 + 2, f"{printhours_2023} hrs\n({printdays_2023:.1f} days)", ha='center', va='bottom', fontsize=10)
plt.text(1, printhours_2024 + 2, f"{printhours_2024} hrs\n({printdays_2024:.1f} days)", ha='center', va='bottom', fontsize=10)

table = plt.table(cellText=[
                    ['2023', f"{printhours_2023}", f"{printdays_2023:.1f}"],
                    ['2024', f"{printhours_2024}", f"{printdays_2024:.1f}"]
                 ],
                 colLabels=['Year', 'Hours', 'Days'],
                 cellLoc='center',
                 bbox=[0.25, -0.35, 0.5, 0.25])
table.scale(1, 1.2)

plt.tight_layout()
# plt.show()

fig = plt.gcf()
save_and_insert(fig, 'total_print_time', doc, caption='')
plt.close(fig)


# === 10: Cost vs Revenue ===


# Data
filament_costs = [3344.26, 5739.30]
parts_costs = [0, 1139.18]
revenues = [2741.58, 3511.94]
losses = [c + p - r for c, p, r in zip(filament_costs, parts_costs, revenues)]

x = [0, 1]
bar_width = 0.25

# Positions for bars
cost_x = [i - bar_width for i in x]
revenue_x = x
loss_x = [i + bar_width for i in x]

plt.figure(figsize=(10, 6))

# Cost bars (stacked)
fc = plt.bar(cost_x, filament_costs, width=bar_width, label='Filament Cost', color='orange')
pc = plt.bar(cost_x, parts_costs, width=bar_width, bottom=filament_costs, label='Parts Cost', color='red')

# Revenue bars
rv = plt.bar(revenue_x, revenues, width=bar_width, label='Revenue', color='green')

# Loss bars
ls = plt.bar(loss_x, losses, width=bar_width, label='Loss', color='grey')

# Add value labels
for i in range(2):
    # Filament
    plt.text(cost_x[i], filament_costs[i] / 2, f"${filament_costs[i]:,.2f}", ha='center', va='center', fontsize=8)
    
    # Parts
    if parts_costs[i] > 0:
        plt.text(cost_x[i], filament_costs[i] + parts_costs[i] / 2, f"${parts_costs[i]:,.2f}", ha='center', va='center', fontsize=8)
    
    # Revenue
    plt.text(revenue_x[i], revenues[i] + 100, f"${revenues[i]:,.2f}", ha='center', va='bottom', fontsize=8, color='green')
    
    # Loss
    plt.text(loss_x[i], losses[i] + 100, f"Loss: ${losses[i]:,.2f}", ha='center', va='bottom', fontsize=8, color='grey')

# Titles and labels
plt.xticks(x, ['2023', '2024'])
plt.title('Cost vs Revenue (with Breakdown)')
plt.ylabel('Dollars ($)')
plt.legend()
plt.tight_layout()

# # Summary annotation
# summary = (
#     "2023:\n"
#     "  Filament Cost: $3,344.26\n"
#     "  Parts Cost: $0.00\n"
#     "  Revenue: $2,741.58\n"
#     "  Loss: $602.68\n\n"
#     "2024:\n"
#     "  Filament Cost: $5,739.30\n"
#     "  Parts Cost: $1,139.18\n"
#     "  Revenue: $3,511.94\n"
#     "  Loss: $3,366.54"
# )

# plt.figtext(0.13, -0.1, summary, fontsize=9, ha='left', va='top', family='monospace')

table_data = [
    ['2023', f"${filament_costs[0]:,.2f}", f"${parts_costs[0]:,.2f}", f"${revenues[0]:,.2f}", f"${losses[0]:,.2f}"],
    ['2024', f"${filament_costs[1]:,.2f}", f"${parts_costs[1]:,.2f}", f"${revenues[1]:,.2f}", f"${losses[1]:,.2f}"]
]
col_labels = ['Year', 'Filament Cost', 'Parts Cost', 'Revenue', 'Loss']

# Add table to plot
table = plt.table(cellText=table_data,
                  colLabels=col_labels,
                  cellLoc='center',
                  bbox=[0.15, -0.35, 0.7, 0.25])
table.scale(1, 1.2)

# plt.show()


fig = plt.gcf()
save_and_insert(fig, 'printing_costing', doc, caption='')
plt.close(fig)


output_doc = "3D_Printer_Report.docx"
doc.save(output_doc)


print(f"Report saved as '{output_doc}'. All plots are embedded.")
